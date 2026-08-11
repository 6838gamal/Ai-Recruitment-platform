"""Resume Parser module services."""

from __future__ import annotations

import re
import time
from pathlib import Path
from typing import Any, Optional

from app.modules.candidates.services import CandidateService


class ResumeParserService:
    """Service responsible for reading and parsing resumes."""

    ALLOWED_EXTENSIONS = {
        ".pdf",
        ".doc",
        ".docx",
    }

    MAX_FILE_SIZE = 10 * 1024 * 1024

    COMMON_SKILLS = {
        "python",
        "javascript",
        "typescript",
        "java",
        "c",
        "c++",
        "c#",
        "php",
        "ruby",
        "go",
        "rust",
        "kotlin",
        "swift",
        "dart",
        "flutter",
        "react",
        "react native",
        "vue",
        "angular",
        "node.js",
        "nodejs",
        "fastapi",
        "flask",
        "django",
        "express",
        "sql",
        "mysql",
        "postgresql",
        "postgres",
        "sqlite",
        "mongodb",
        "redis",
        "docker",
        "kubernetes",
        "aws",
        "azure",
        "gcp",
        "git",
        "github",
        "gitlab",
        "linux",
        "html",
        "css",
        "tailwind",
        "machine learning",
        "deep learning",
        "artificial intelligence",
        "ai",
        "data science",
        "pandas",
        "numpy",
        "tensorflow",
        "pytorch",
        "scikit-learn",
        "openai",
        "rest api",
        "graphql",
        "microservices",
        "agile",
        "scrum",
    }

    def __init__(self, db=None):
        self.db = db

    # ================================================================
    # Main parser
    # ================================================================

    def parse_file(
        self,
        file_path: str | Path,
        filename: Optional[str] = None,
    ) -> dict[str, Any]:
        """
        Parse a resume file and return structured data.
        """

        started_at = time.perf_counter()

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(
                f"Resume file not found: {file_path}"
            )

        file_size = file_path.stat().st_size

        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(
                "Resume file exceeds the maximum size of 10MB."
            )

        original_filename = (
            filename
            or file_path.name
        )

        extension = file_path.suffix.lower()

        if extension not in self.ALLOWED_EXTENSIONS:
            raise ValueError(
                "Unsupported resume format. "
                "Supported formats: PDF, DOC, DOCX."
            )

        text = self.extract_text(
            file_path=file_path,
            extension=extension,
        )

        text = self.clean_text(text)

        if not text.strip():
            raise ValueError(
                "Could not extract readable text from the resume."
            )

        parsed = self.extract_information(text)

        parsed["filename"] = original_filename
        parsed["file_path"] = str(file_path)
        parsed["file_type"] = extension.lstrip(".")
        parsed["file_size"] = file_size
        parsed["resume_text"] = text

        parsed["parse_time"] = round(
            time.perf_counter() - started_at,
            3,
        )

        return parsed

    # ================================================================
    # Text extraction
    # ================================================================

    def extract_text(
        self,
        file_path: Path,
        extension: str,
    ) -> str:
        """Extract text from PDF, DOCX or DOC."""

        if extension == ".pdf":
            return self._extract_pdf_text(
                file_path
            )

        if extension == ".docx":
            return self._extract_docx_text(
                file_path
            )

        if extension == ".doc":
            return self._extract_doc_text(
                file_path
            )

        raise ValueError(
            f"Unsupported file extension: {extension}"
        )

    @staticmethod
    def _extract_pdf_text(
        file_path: Path,
    ) -> str:
        """Extract text from PDF."""

        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError(
                "pypdf is required to parse PDF files. "
                "Install it with: pip install pypdf"
            ) from exc

        reader = PdfReader(str(file_path))

        pages = []

        for page in reader.pages:

            try:
                page_text = page.extract_text()
            except Exception:
                page_text = ""

            if page_text:
                pages.append(page_text)

        return "\n".join(pages)

    @staticmethod
    def _extract_docx_text(
        file_path: Path,
    ) -> str:
        """Extract text from DOCX."""

        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required to parse DOCX files. "
                "Install it with: pip install python-docx"
            ) from exc

        document = Document(
            str(file_path)
        )

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        # Also extract text from tables.
        for table in document.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    value = cell.text.strip()

                    if value:
                        cells.append(value)

                if cells:
                    paragraphs.append(
                        " | ".join(cells)
                    )

        return "\n".join(paragraphs)

    @staticmethod
    def _extract_doc_text(
        file_path: Path,
    ) -> str:
        """
        Extract text from legacy .DOC files.

        Uses antiword when available.
        """

        import subprocess

        try:
            result = subprocess.run(
                [
                    "antiword",
                    str(file_path),
                ],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )

            if result.returncode == 0:
                return result.stdout

        except (
            FileNotFoundError,
            subprocess.SubprocessError,
        ):
            pass

        raise RuntimeError(
            "Legacy DOC files require 'antiword'. "
            "Install antiword on the server or convert "
            "the document to DOCX/PDF."
        )

    # ================================================================
    # Text cleaning
    # ================================================================

    @staticmethod
    def clean_text(
        text: str,
    ) -> str:
        """Clean extracted resume text."""

        if not text:
            return ""

        text = text.replace(
            "\x00",
            " ",
        )

        text = re.sub(
            r"[ \t]+",
            " ",
            text,
        )

        text = re.sub(
            r"\n{3,}",
            "\n\n",
            text,
        )

        return text.strip()

    # ================================================================
    # Information extraction
    # ================================================================

    def extract_information(
        self,
        text: str,
    ) -> dict[str, Any]:
        """Extract structured information from resume text."""

        email = self.extract_email(text)

        phone = self.extract_phone(text)

        name = self.extract_name(
            text,
            email=email,
        )

        title = self.extract_title(text)

        location = self.extract_location(text)

        skills = self.extract_skills(text)

        experience = self.extract_experience(text)

        education = self.extract_education(text)

        years_of_experience = (
            self.calculate_years_of_experience(
                experience
            )
        )

        summary = self.extract_summary(
            text
        )

        return {
            "first_name": name.get(
                "first_name"
            ),
            "last_name": name.get(
                "last_name"
            ),
            "full_name": name.get(
                "full_name"
            ),
            "email": email,
            "phone": phone,
            "location": location,
            "title": title,
            "job_title": title,
            "summary": summary,
            "skills": skills,
            "experience": experience,
            "education": education,
            "years_of_experience": years_of_experience,
        }

    # ================================================================
    # Email
    # ================================================================

    @staticmethod
    def extract_email(
        text: str,
    ) -> Optional[str]:
        """Extract the first email address."""

        match = re.search(
            r"\b[A-Za-z0-9._%+-]+"
            r"@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
            text,
        )

        if not match:
            return None

        return match.group(0).strip()

    # ================================================================
    # Phone
    # ================================================================

    @staticmethod
    def extract_phone(
        text: str,
    ) -> Optional[str]:
        """Extract a likely phone number."""

        patterns = [
            r"\+?\d[\d\s().-]{7,}\d",
        ]

        for pattern in patterns:

            matches = re.findall(
                pattern,
                text,
            )

            for value in matches:

                value = value.strip()

                digits = re.sub(
                    r"\D",
                    "",
                    value,
                )

                if 8 <= len(digits) <= 15:
                    return value

        return None

    # ================================================================
    # Name
    # ================================================================

    def extract_name(
        self,
        text: str,
        email: Optional[str] = None,
    ) -> dict[str, Optional[str]]:
        """Extract candidate name."""

        lines = [
            line.strip()
            for line in text.splitlines()
            if line.strip()
        ]

        ignored = {
            "resume",
            "cv",
            "curriculum vitae",
            "profile",
            "personal information",
            "contact",
            "contact information",
            "experience",
            "work experience",
            "education",
            "skills",
            "professional summary",
        }

        for line in lines[:15]:

            clean = line.strip()

            lower = clean.lower()

            if lower in ignored:
                continue

            if email and email.lower() in lower:
                continue

            if re.search(
                r"\d{4}",
                clean,
            ):
                continue

            if "@" in clean:
                continue

            if re.search(
                r"\+?\d[\d\s().-]{7,}",
                clean,
            ):
                continue

            words = clean.split()

            if 2 <= len(words) <= 5:

                valid = True

                for word in words:

                    if not re.match(
                        r"^[A-Za-z\u0600-\u06FF'’-]+$",
                        word,
                    ):
                        valid = False
                        break

                if valid:

                    return {
                        "full_name": clean,
                        "first_name": words[0],
                        "last_name": (
                            " ".join(words[1:])
                            if len(words) > 1
                            else ""
                        ),
                    }

        return {
            "full_name": None,
            "first_name": None,
            "last_name": None,
        }

    # ================================================================
    # Job title
    # ================================================================

    def extract_title(
        self,
        text: str,
    ) -> Optional[str]:
        """Extract a likely professional title."""

        title_patterns = [
            r"(?:job title|title|position)\s*[:\-]\s*(.+)",
            r"(?:role)\s*[:\-]\s*(.+)",
            r"(?:profession)\s*[:\-]\s*(.+)",
        ]

        for pattern in title_patterns:

            match = re.search(
                pattern,
                text,
                re.IGNORECASE,
            )

            if match:

                value = match.group(1).strip()

                if value:
                    return value[:255]

        common_titles = [
            "software engineer",
            "software developer",
            "senior software engineer",
            "backend developer",
            "frontend developer",
            "full stack developer",
            "full-stack developer",
            "mobile developer",
            "flutter developer",
            "python developer",
            "java developer",
            "php developer",
            "data scientist",
            "data analyst",
            "machine learning engineer",
            "ai engineer",
            "devops engineer",
            "cloud engineer",
            "project manager",
            "product manager",
            "hr manager",
            "human resources manager",
            "accountant",
            "marketing manager",
            "sales manager",
            "designer",
            "ui ux designer",
        ]

        lower_text = text.lower()

        for title in common_titles:

            if title in lower_text:
                return title.title()

        return None

    # ================================================================
    # Location
    # ================================================================

    @staticmethod
    def extract_location(
        text: str,
    ) -> Optional[str]:
        """Extract a likely location."""

        patterns = [
            r"(?:location|address|based in|residence)"
            r"\s*[:\-]\s*(.+)",
        ]

        for pattern in patterns:

            match = re.search(
                pattern,
                text,
                re.IGNORECASE,
            )

            if match:

                value = match.group(1).strip()

                value = value.splitlines()[0]

                if value:
                    return value[:255]

        return None

    # ================================================================
    # Skills
    # ================================================================

    def extract_skills(
        self,
        text: str,
    ) -> list[str]:
        """Extract known technical and professional skills."""

        lower_text = text.lower()

        found: list[str] = []

        for skill in sorted(
            self.COMMON_SKILLS,
            key=len,
            reverse=True,
        ):

            pattern = (
                r"(?<![a-z0-9])"
                + re.escape(skill.lower())
                + r"(?![a-z0-9])"
            )

            if re.search(
                pattern,
                lower_text,
            ):
                found.append(skill)

        return sorted(
            set(found),
            key=str.lower,
        )

    # ================================================================
    # Experience
    # ================================================================

    def extract_experience(
        self,
        text: str,
    ) -> list[dict[str, Any]]:
        """Extract basic work experience entries."""

        section = self.extract_section(
            text,
            [
                "experience",
                "work experience",
                "employment",
                "professional experience",
                "career history",
            ],
        )

        if not section:
            return []

        lines = [
            line.strip()
            for line in section.splitlines()
            if line.strip()
        ]

        entries: list[dict[str, Any]] = []

        current: Optional[dict[str, Any]] = None

        date_pattern = re.compile(
            r"("
            r"(?:19|20)\d{2}"
            r")"
            r"\s*[-–]\s*"
            r"("
            r"(?:19|20)\d{2}"
            r"|present"
            r"|current"
            r")",
            re.IGNORECASE,
        )

        for line in lines:

            date_match = date_pattern.search(
                line
            )

            if date_match:

                if current:
                    entries.append(current)

                before_date = (
                    line[:date_match.start()]
                    .strip(" -–|")
                )

                current = {
                    "position": before_date
                    or None,
                    "company": None,
                    "start_date": date_match.group(1),
                    "end_date": date_match.group(2),
                    "description": "",
                }

                continue

            if current is None:

                current = {
                    "position": line,
                    "company": None,
                    "start_date": None,
                    "end_date": None,
                    "description": "",
                }

                continue

            if not current["company"]:

                current["company"] = line

                continue

            if current["description"]:

                current["description"] += " " + line

            else:

                current["description"] = line

        if current:
            entries.append(current)

        return entries[:20]

    # ================================================================
    # Education
    # ================================================================

    def extract_education(
        self,
        text: str,
    ) -> list[dict[str, Any]]:
        """Extract basic education entries."""

        section = self.extract_section(
            text,
            [
                "education",
                "academic background",
                "academic qualifications",
            ],
        )

        if not section:
            return []

        lines = [
            line.strip()
            for line in section.splitlines()
            if line.strip()
        ]

        entries: list[dict[str, Any]] = []

        for line in lines[:20]:

            year_match = re.search(
                r"\b(19|20)\d{2}\b",
                line,
            )

            year = (
                year_match.group(0)
                if year_match
                else None
            )

            value = re.sub(
                r"\b(19|20)\d{2}\b",
                "",
                line,
            ).strip(" -–|")

            if not value:
                continue

            degree = value
            institution = None

            separators = [
                " - ",
                " | ",
                " at ",
                " from ",
            ]

            for separator in separators:

                if separator in value:

                    parts = value.split(
                        separator,
                        1,
                    )

                    degree = parts[0].strip()
                    institution = parts[1].strip()

                    break

            entries.append(
                {
                    "degree": degree,
                    "institution": institution,
                    "year": year,
                }
            )

        return entries

    # ================================================================
    # Summary
    # ================================================================

    def extract_summary(
        self,
        text: str,
    ) -> Optional[str]:
        """Extract professional summary."""

        section = self.extract_section(
            text,
            [
                "summary",
                "professional summary",
                "profile",
                "objective",
                "career objective",
                "about me",
            ],
        )

        if not section:
            return None

        section = self.clean_text(
            section
        )

        if not section:
            return None

        return section[:3000]

    # ================================================================
    # Sections
    # ================================================================

    @staticmethod
    def extract_section(
        text: str,
        section_names: list[str],
    ) -> str:
        """Extract text belonging to a resume section."""

        lines = text.splitlines()

        start_index = None

        normalized_names = {
            name.lower().strip()
            for name in section_names
        }

        for index, line in enumerate(lines):

            clean = line.strip().lower()

            clean = re.sub(
                r"[:\-]+$",
                "",
                clean,
            )

            if clean in normalized_names:

                start_index = index + 1
                break

        if start_index is None:
            return ""

        section_lines: list[str] = []

        common_end_sections = {
            "experience",
            "work experience",
            "employment",
            "education",
            "skills",
            "technical skills",
            "certifications",
            "languages",
            "projects",
            "references",
            "summary",
            "professional summary",
            "profile",
            "objective",
        }

        for line in lines[start_index:]:

            clean = line.strip().lower()

            clean = re.sub(
                r"[:\-]+$",
                "",
                clean,
            )

            if (
                clean in common_end_sections
                and section_lines
            ):
                break

            section_lines.append(line)

        return "\n".join(
            section_lines
        ).strip()

    # ================================================================
    # Years of experience
    # ================================================================

    @staticmethod
    def calculate_years_of_experience(
        experience: list[dict[str, Any]],
    ) -> float:
        """Calculate approximate total experience."""

        total_months = 0

        current_year = time.localtime().tm_year

        for item in experience:

            start = item.get(
                "start_date"
            )

            end = item.get(
                "end_date"
            )

            if not start:
                continue

            start_match = re.search(
                r"(19|20)\d{2}",
                str(start),
            )

            if not start_match:
                continue

            start_year = int(
                start_match.group(0)
            )

            if not end or str(end).lower() in {
                "present",
                "current",
            }:
                end_year = current_year

            else:

                end_match = re.search(
                    r"(19|20)\d{2}",
                    str(end),
                )

                if not end_match:
                    continue

                end_year = int(
                    end_match.group(0)
                )

            if end_year >= start_year:

                total_months += (
                    end_year - start_year
                ) * 12

        return round(
            total_months / 12,
            1,
        )

    # ================================================================
    # Matching
    # ================================================================

    def match_with_candidates(
        self,
        parsed_resume: dict[str, Any],
        *,
        limit: int = 20,
    ) -> list[dict[str, Any]]:
        """Compare parsed resume with existing candidates."""

        if self.db is None:
            return []

        service = CandidateService(
            self.db
        )

        return service.match_resume(
            parsed_resume,
            limit=limit,
        )

    # ================================================================
    # Complete parsing + matching
    # ================================================================

    def parse_and_match(
        self,
        file_path: str | Path,
        filename: Optional[str] = None,
        *,
        limit: int = 20,
    ) -> dict[str, Any]:
        """Parse a resume and compare it with candidates."""

        parsed_resume = self.parse_file(
            file_path=file_path,
            filename=filename,
        )

        matches = self.match_with_candidates(
            parsed_resume,
            limit=limit,
        )

        parsed_resume["matches"] = matches

        parsed_resume["best_match_score"] = (
            matches[0]["match_score"]
            if matches
            else 0
        )

        return parsed_resume
